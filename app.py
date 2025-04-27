#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
MinerU Web UI 应用程序

这个应用程序提供了一个基于Gradio的Web UI界面，用于将PDF、图片、DOCX等文档转换为Markdown格式。
支持预览Markdown（包括图片预览）和下载转换后的文件。
"""

import os
import time
import shutil
import zipfile
import tempfile
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed

import gradio as gr
import docx
from PIL import Image
from loguru import logger

# 导入MinerU相关模块
import sys
sys.path.append('..')
from magic_pdf.data.data_reader_writer import FileBasedDataWriter, FileBasedDataReader
from magic_pdf.data.dataset import PymuDocDataset
from magic_pdf.model.doc_analyze_by_custom_model import doc_analyze
from magic_pdf.config.enums import SupportedPdfParseMethod

# 导入自定义工具类
from utils import ParallelProcessor, FileUtils, ProgressTracker


class MinerUWebConverter:
    """MinerU文档转换Web界面控制器类"""
    
    def __init__(self, output_dir="output"):
        """
        初始化转换器
        
        Args:
            output_dir: 输出目录
        """
        self.output_dir = output_dir
        FileUtils.ensure_dir(output_dir)
        self.processor = ParallelProcessor()
        logger.info(f"初始化MinerU Web转换器，输出目录: {output_dir}")
        
    def _process_pdf(self, file_path, task_id):
        """
        处理PDF文件
        
        Args:
            file_path: PDF文件路径
            task_id: 任务ID
            
        Returns:
            转换结果信息，包含Markdown文本和图片路径
        """
        logger.info(f"处理PDF文件: {file_path}")
        
        # 创建临时输出目录
        name_without_extension = os.path.basename(file_path).split('.')[0]
        task_output_dir = os.path.join(self.output_dir, task_id)
        local_image_dir = os.path.join(task_output_dir, "images")
        FileUtils.ensure_dir(local_image_dir)
        
        image_writer = FileBasedDataWriter(local_image_dir)
        md_writer = FileBasedDataWriter(task_output_dir)
        
        # 读取PDF内容
        reader = FileBasedDataReader("")
        pdf_bytes = reader.read(file_path)
        
        # 创建数据集实例
        ds = PymuDocDataset(pdf_bytes)
        
        # 根据文档类型选择处理方式
        if ds.classify() == SupportedPdfParseMethod.OCR:
            infer_result = ds.apply(doc_analyze, ocr=True)
            pipe_result = infer_result.pipe_ocr_mode(image_writer)
        else:
            infer_result = ds.apply(doc_analyze, ocr=False)
            pipe_result = infer_result.pipe_txt_mode(image_writer)
        
        # 获取Markdown内容
        image_dir = "images"
        md_content = pipe_result.get_markdown(image_dir)
        
        # 保存Markdown文件
        md_file_path = os.path.join(task_output_dir, f"{name_without_extension}.md")
        pipe_result.dump_md(md_writer, f"{name_without_extension}.md", image_dir)
        
        # 创建ZIP压缩包
        zip_path = os.path.join(task_output_dir, f"{name_without_extension}.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            # 添加Markdown文件
            zipf.write(md_file_path, os.path.basename(md_file_path))
            
            # 添加图片文件夹
            for img_file in os.listdir(local_image_dir):
                img_path = os.path.join(local_image_dir, img_file)
                if os.path.isfile(img_path):
                    zipf.write(img_path, os.path.join("images", img_file))
        
        return {
            "md_content": md_content,
            "zip_path": zip_path,
            "md_file_path": md_file_path,
            "image_dir": local_image_dir
        }
    
    def _process_docx(self, file_path, task_id):
        """
        处理DOCX文件
        
        Args:
            file_path: DOCX文件路径
            task_id: 任务ID
            
        Returns:
            转换结果信息
        """
        logger.info(f"处理DOCX文件: {file_path}")
        
        # 创建临时输出目录
        name_without_extension = os.path.basename(file_path).split('.')[0]
        task_output_dir = os.path.join(self.output_dir, task_id)
        local_image_dir = os.path.join(task_output_dir, "images")
        FileUtils.ensure_dir(local_image_dir)
        
        # 读取DOCX文件
        doc = docx.Document(file_path)
        
        # 提取文本和图片
        md_content = []
        image_index = 1
        
        # 处理段落和图片
        for para in doc.paragraphs:
            # 处理段落文本
            if para.text.strip():
                # 获取段落级别（如果是标题）
                if para.style.name.startswith('Heading'):
                    heading_level = int(para.style.name.replace('Heading', ''))
                    md_content.append(f"{'#' * heading_level} {para.text}")
                else:
                    md_content.append(para.text)
            
            # 处理段落中的图片（如果有）
            for run in para.runs:
                for shape in run._element.drawing_lst:
                    # 保存图片
                    image_path = os.path.join(local_image_dir, f"image_{image_index}.png")
                    
                    # 提取并保存图片
                    if self._extract_and_save_docx_image(shape, image_path):
                        md_content.append(f"![图片 {image_index}](images/image_{image_index}.png)")
                        image_index += 1
        
        # 处理表格
        for table in doc.tables:
            table_md = []
            # 获取表头
            header_row = table.rows[0]
            header_cells = [cell.text.strip() for cell in header_row.cells]
            
            # 表格标记开始
            table_md.append('| ' + ' | '.join(header_cells) + ' |')
            table_md.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')
            
            # 获取表格数据行
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                table_md.append('| ' + ' | '.join(cells) + ' |')
            
            md_content.append('\n'.join(table_md))
        
        # 合并Markdown内容
        full_md_content = '\n\n'.join(md_content)
        
        # 保存Markdown文件
        md_file_path = os.path.join(task_output_dir, f"{name_without_extension}.md")
        with open(md_file_path, 'w', encoding='utf-8') as f:
            f.write(full_md_content)
        
        # 创建ZIP压缩包
        zip_path = os.path.join(task_output_dir, f"{name_without_extension}.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            # 添加Markdown文件
            zipf.write(md_file_path, os.path.basename(md_file_path))
            
            # 添加图片文件夹
            for img_file in os.listdir(local_image_dir):
                img_path = os.path.join(local_image_dir, img_file)
                if os.path.isfile(img_path):
                    zipf.write(img_path, os.path.join("images", img_file))
        
        return {
            "md_content": full_md_content,
            "zip_path": zip_path,
            "md_file_path": md_file_path,
            "image_dir": local_image_dir
        }
    
    def _extract_and_save_docx_image(self, shape, image_path):
        """
        从DOCX文件中提取并保存图片
        
        Args:
            shape: 文档中的图形对象
            image_path: 保存图片的路径
            
        Returns:
            布尔值，表示是否成功保存图片
        """
        try:
            # 提取图片并保存
            with open(image_path, 'wb') as f:
                f.write(shape.blob)
            return True
        except Exception as e:
            logger.error(f"保存图片失败: {e}")
            return False
    
    def _process_image(self, file_path, task_id):
        """
        处理图片文件
        
        Args:
            file_path: 图片文件路径
            task_id: 任务ID
            
        Returns:
            转换结果信息
        """
        logger.info(f"处理图片文件: {file_path}")
        
        # 创建临时输出目录
        name_without_extension = os.path.basename(file_path).split('.')[0]
        task_output_dir = os.path.join(self.output_dir, task_id)
        local_image_dir = os.path.join(task_output_dir, "images")
        FileUtils.ensure_dir(local_image_dir)
        
        # 复制图片到输出目录
        output_image_path = os.path.join(local_image_dir, os.path.basename(file_path))
        shutil.copy(file_path, output_image_path)
        
        # 使用MinerU进行OCR识别
        img = Image.open(file_path)
        img_data = None
        with open(file_path, 'rb') as f:
            img_data = f.read()
        
        # 创建数据集实例
        ds = PymuDocDataset(img_data)
        
        # 执行OCR
        infer_result = ds.apply(doc_analyze, ocr=True)
        
        # 设置文件写入器
        image_writer = FileBasedDataWriter(local_image_dir)
        md_writer = FileBasedDataWriter(task_output_dir)
        
        # 获取处理结果
        pipe_result = infer_result.pipe_ocr_mode(image_writer)
        image_dir = "images"
        md_content = pipe_result.get_markdown(image_dir)
        
        # 如果OCR没有提取文本，添加图片Markdown
        if not md_content.strip():
            md_content = f"![{name_without_extension}](images/{os.path.basename(file_path)})"
        
        # 保存Markdown文件
        md_file_path = os.path.join(task_output_dir, f"{name_without_extension}.md")
        with open(md_file_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        # 创建ZIP压缩包
        zip_path = os.path.join(task_output_dir, f"{name_without_extension}.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            # 添加Markdown文件
            zipf.write(md_file_path, os.path.basename(md_file_path))
            
            # 添加图片文件
            zipf.write(output_image_path, os.path.join("images", os.path.basename(file_path)))
        
        return {
            "md_content": md_content,
            "zip_path": zip_path,
            "md_file_path": md_file_path,
            "image_dir": local_image_dir
        }
    
    def process_file(self, file_path):
        """
        处理上传的文件
        
        Args:
            file_path: 上传文件的路径
            
        Returns:
            处理结果，包括Markdown预览和下载链接
        """
        # 生成唯一任务ID
        task_id = f"task_{int(time.time())}"
        
        # 根据文件类型执行不同的处理
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            result = self._process_pdf(file_path, task_id)
        elif file_ext in ['.doc', '.docx']:
            result = self._process_docx(file_path, task_id)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
            result = self._process_image(file_path, task_id)
        else:
            return "不支持的文件类型", None
        
        return result["md_content"], result["zip_path"]
    
    def batch_process_files(self, file_paths):
        """
        批量处理多个文件
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            包含所有处理结果的ZIP文件路径
        """
        # 生成批处理任务ID
        task_id = f"batch_{int(time.time())}"
        task_output_dir = os.path.join(self.output_dir, task_id)
        FileUtils.ensure_dir(task_output_dir)
        
        # 创建进度跟踪器
        progress = ProgressTracker(len(file_paths), "批量处理文件")
        
        # 定义处理回调函数
        def process_callback(file_path, result):
            progress.update()
            logger.info(f"已处理文件: {file_path}, 进度: {progress.get_progress()['percentage']:.2f}%")
        
        # 使用并行处理器处理文件
        results = self.processor.process_items(
            file_paths,
            self.process_file,
            process_callback
        )
        
        # 创建汇总ZIP包
        batch_zip_path = os.path.join(task_output_dir, "batch_results.zip")
        with zipfile.ZipFile(batch_zip_path, 'w') as batch_zip:
            for i, result in enumerate(results):
                if isinstance(result, tuple) and len(result) == 2 and result[1]:
                    md_content, zip_path = result
                    if os.path.exists(zip_path):
                        # 解压单个文件的ZIP到临时目录
                        with tempfile.TemporaryDirectory() as temp_dir:
                            with zipfile.ZipFile(zip_path, 'r') as file_zip:
                                file_zip.extractall(temp_dir)
                            
                            # 将文件添加到汇总ZIP，使用文件名作为子目录
                            base_name = os.path.basename(zip_path).split('.')[0]
                            for root, _, files in os.walk(temp_dir):
                                for file in files:
                                    file_path = os.path.join(root, file)
                                    rel_path = os.path.relpath(file_path, temp_dir)
                                    batch_zip.write(file_path, os.path.join(base_name, rel_path))
        
        return batch_zip_path


def create_ui():
    """创建Gradio Web UI界面"""
    
    # 初始化转换器
    converter = MinerUWebConverter()
    
    with gr.Blocks(title="MinerU 文档转换工具") as app:
        gr.Markdown("""
        # MinerU 文档转换工具
        
        将PDF、DOCX和图片文件转换为Markdown格式。支持图片预览和下载转换后的文件。
        """)
        
        with gr.Tab("单文件转换"):
            with gr.Row():
                file_input = gr.File(label="上传文件（支持 PDF、DOCX、JPG、PNG）")
            
            with gr.Row():
                convert_btn = gr.Button("转换为Markdown", variant="primary")
            
            with gr.Row():
                with gr.Column():
                    md_preview = gr.Markdown(label="Markdown预览")
                with gr.Column():
                    output_file = gr.File(label="下载ZIP文件")
        
        with gr.Tab("批量转换"):
            with gr.Row():
                files_input = gr.Files(label="上传多个文件（支持 PDF、DOCX、JPG、PNG）")
            
            with gr.Row():
                batch_convert_btn = gr.Button("批量转换", variant="primary")
                
            with gr.Row():
                processing_status = gr.Textbox(label="处理状态", value="就绪")
            
            with gr.Row():
                batch_output = gr.File(label="下载汇总ZIP文件")
        
        # 单文件转换事件
        convert_btn.click(
            fn=converter.process_file,
            inputs=[file_input],
            outputs=[md_preview, output_file]
        )
        
        # 批量转换事件
        def batch_process(files):
            result = converter.batch_process_files(files)
            return result, "处理完成"
        
        batch_convert_btn.click(
            fn=batch_process,
            inputs=[files_input],
            outputs=[batch_output, processing_status]
        )
    
    return app


if __name__ == "__main__":
    # 设置日志
    logger.add("mineru_web.log", rotation="10 MB")
    
    # 创建并启动Gradio应用
    app = create_ui()
    app.launch(server_name="0.0.0.0", server_port=7860, share=True) 